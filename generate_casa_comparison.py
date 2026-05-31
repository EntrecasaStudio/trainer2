#!/usr/bin/env python3
"""Generate DOCX with comparative tables of all CASA routines."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from collections import Counter

# ── Data ──────────────────────────────────────────────────────────────────────

LEAN_PRESS = [
    {
        'name': 'Casa Press A — Lean',
        'circuits': [
            ('PIERNAS·ESPALDA', ['Sentadilla goblet', 'Zancadas con chaleco de peso', 'TRX row']),
            ('PECHO', ['Floor press', 'Flexiones con chaleco', 'TRX chest press']),
            ('HOMBROS·TRÍCEPS', ['Press militar', 'Fondos en banco con chaleco', 'Vuelos laterales']),
            ('CORE·BÍCEPS', ['Ab wheel', 'Plancha con chaleco', 'Bíceps curl']),
            ('HIIT', ['Mountain climbers', 'Burpees', 'Jumping jacks']),
        ]
    },
    {
        'name': 'Casa Press B — Lean',
        'circuits': [
            ('PIERNAS·ESPALDA', ['Sentadilla sumo', 'Remo', 'Sentadilla búlgara']),
            ('PECHO', ['Flexiones explosivas', 'TRX archer press', 'TRX chest press']),
            ('HOMBROS·TRÍCEPS', ['Arnold press con kettlebell', 'Extensión de tríceps con kettlebell', 'Elevaciones de hombro adelante']),
            ('CORE·BÍCEPS', ['Abs complex', 'Plancha con elevación alternada', 'Bíceps curl martillo']),
            ('HIIT', ['Bear crawl', 'Tuck jumps', 'Saltos laterales']),
        ]
    },
    {
        'name': 'Casa Press C — Lean',
        'circuits': [
            ('PIERNAS·ESPALDA', ['Sentadilla con salto', 'Zancadas con kettlebell', 'TRX power pull']),
            ('PECHO', ['Flexiones diamante', 'Floor press', 'Flexiones con chaleco']),
            ('HOMBROS·TRÍCEPS', ['Fondos en banco', 'Vuelos laterales en equilibrio con kettlebell', 'Banda triceps pushdown']),
            ('CORE·BÍCEPS', ['Plancha estrella con peso', 'Dead bug', 'Bíceps curl en TRX']),
            ('HIIT', ['Sentadilla con salto', 'Estocada con salto', 'Mountain climbers']),
        ]
    },
    {
        'name': 'Casa Press D — Lean',
        'circuits': [
            ('PIERNAS·ESPALDA', ['Peso muerto a una pierna', 'Sentadilla sumo', 'Remo alto en TRX']),
            ('PECHO', ['TRX chest press', 'Flexiones inclinadas', 'Flexiones explosivas']),
            ('HOMBROS·TRÍCEPS', ['Tríceps alto en TRX', 'Arnold press con kettlebell', 'Extensión de tríceps con banda']),
            ('CORE·BÍCEPS', ['Plancha lateral con chaleco', 'Abs complex', 'Bíceps curl en equilibrio con kettlebell']),
            ('HIIT', ['Bear crawl', 'Burpees', 'Saltos laterales']),
        ]
    },
    {
        'name': 'Casa Press E — Lean',
        'circuits': [
            ('PIERNAS·ESPALDA', ['Sentadilla goblet', 'Step up', 'TRX face pull']),
            ('PECHO', ['Flexiones diamante', 'TRX archer press', 'TRX chest press']),
            ('HOMBROS·TRÍCEPS', ['Press militar', 'Fondos en banco con chaleco', 'Vuelos laterales en equilibrio con kettlebell']),
            ('CORE·BÍCEPS', ['Ab wheel', 'Caminata a plancha', 'Bíceps curl']),
            ('HIIT', ['Tuck jumps', 'Estocada con salto', 'Jumping jacks']),
        ]
    },
    {
        'name': 'Casa Press F — Lean',
        'circuits': [
            ('PIERNAS·ESPALDA', ['Peso muerto a una pierna', 'Sentadilla búlgara', 'Banda pull-apart']),
            ('PECHO', ['Flexiones inclinadas', 'TRX chest press', 'Floor press']),
            ('HOMBROS·TRÍCEPS', ['Extensión de tríceps con banda', 'Elevaciones de hombro adelante', 'Tríceps alto en TRX']),
            ('CORE·BÍCEPS', ['Hollow body', 'Plancha con elevación alternada', 'Bíceps curl martillo']),
            ('HIIT', ['Saltos laterales', 'Sentadilla con salto', 'Bear crawl']),
        ]
    },
]

LEAN_PULL = [
    {
        'name': 'Casa Pull A — Lean',
        'circuits': [
            ('PIERNAS·CORE', ['Empuje de cadera', 'Calf raises', 'Plancha con chaleco']),
            ('ESPALDA', ['Dominadas australianas', 'Remo alto en TRX', 'TRX face pull']),
            ('BÍCEPS·CORE', ['Bíceps curl en TRX', 'Dead bug', 'Bíceps curl']),
            ('PECHO·TRÍCEPS', ['Fondos en banco', 'Flexiones', 'Extensión de tríceps con banda']),
            ('HIIT', ['Sentadilla con salto', 'Bear crawl', 'Saltos laterales']),
        ]
    },
    {
        'name': 'Casa Pull B — Lean',
        'circuits': [
            ('PIERNAS·CORE', ['Sentadilla goblet', 'Plancha lateral con chaleco', 'Squat to press con kettlebell']),
            ('ESPALDA', ['TRX row', 'TRX power pull', 'Banda pull-apart']),
            ('BÍCEPS·CORE', ['Bíceps curl martillo', 'Pallof press', 'Bíceps curl concentrado']),
            ('PECHO·TRÍCEPS', ['Tríceps alto en TRX', 'Flexiones inclinadas', 'Banda triceps pushdown']),
            ('HIIT', ['Burpees', 'Jumping jacks', 'Estocada con salto']),
        ]
    },
    {
        'name': 'Casa Pull C — Lean',
        'circuits': [
            ('PIERNAS·CORE', ['Zancadas con chaleco de peso', 'Hollow body', 'Sentadilla búlgara con TRX']),
            ('ESPALDA', ['Remo alto en TRX', 'Dominadas australianas', 'TRX face pull']),
            ('BÍCEPS·CORE', ['Bíceps curl en equilibrio con kettlebell', 'Abs complex', 'Bíceps curl en TRX']),
            ('PECHO·TRÍCEPS', ['Fondos en banco con chaleco', 'Flexiones diamante', 'Extensión de tríceps con kettlebell']),
            ('HIIT', ['Mountain climbers', 'Tuck jumps', 'Jumping jacks']),
        ]
    },
    {
        'name': 'Casa Pull D — Lean',
        'circuits': [
            ('PIERNAS·CORE', ['Sentadilla búlgara', 'Empuje de cadera', 'Plancha toque de hombro']),
            ('ESPALDA', ['TRX row', 'Remo', 'Remo alto en TRX']),
            ('BÍCEPS·CORE', ['Bíceps curl', 'Bird-dog', 'Bíceps curl martillo']),
            ('PECHO·TRÍCEPS', ['Fondos en banco', 'TRX chest press', 'Extensión de tríceps con banda']),
            ('HIIT', ['Sentadilla con salto', 'Burpees', 'Saltos laterales']),
        ]
    },
    {
        'name': 'Casa Pull E — Lean',
        'circuits': [
            ('PIERNAS·CORE', ['Empuje de cadera', 'Peso muerto a una pierna', 'Copenhague']),
            ('ESPALDA', ['Dominadas australianas', 'Remo alto en TRX', 'TRX face pull']),
            ('BÍCEPS·CORE', ['Bíceps curl en TRX', 'Plancha estrella con peso', 'Bíceps curl en equilibrio con kettlebell']),
            ('PECHO·TRÍCEPS', ['Fondos en banco con chaleco', 'Flexiones', 'Extensión de tríceps con kettlebell']),
            ('HIIT', ['Mountain climbers', 'Bear crawl', 'Jumping jacks']),
        ]
    },
    {
        'name': 'Casa Pull F — Lean',
        'circuits': [
            ('PIERNAS·CORE', ['Step up', 'Sentadilla sumo', 'Abs complex']),
            ('ESPALDA', ['TRX row', 'TRX power pull', 'Remo']),
            ('BÍCEPS·CORE', ['Bíceps curl en equilibrio con kettlebell', 'Bird-dog', 'Bíceps curl concentrado']),
            ('PECHO·TRÍCEPS', ['Flexiones explosivas', 'Tríceps alto en TRX', 'TRX chest press']),
            ('HIIT', ['Caminata a plancha', 'Tuck jumps', 'Estocada con salto']),
        ]
    },
]

NAT_PRESS = [
    {
        'name': 'Casa Press A — Nat',
        'circuits': [
            ('PIERNAS·GLÚTEOS', ['Sumo squat to calf raise', 'Hip thrust a una pierna', 'Sentadilla goblet']),
            ('PECHO·HOMBROS', ['Floor press', 'Press militar', 'Flexiones inclinadas']),
            ('HOMBROS·TRÍCEPS', ['Fondos en banco', 'Vuelos laterales', 'Extensión de tríceps con banda']),
            ('CORE·PIERNAS', ['Ab wheel', 'Plancha con chaleco', 'Empuje de cadera en piso con peso']),
            ('HIIT', ['Jumping jacks', 'Mountain climbers', 'Burpees']),
        ]
    },
    {
        'name': 'Casa Press B — Nat',
        'circuits': [
            ('PIERNAS·GLÚTEOS', ['Elevated side leg lifts', 'Déficit reverse lunge', 'Step up']),
            ('PECHO·HOMBROS', ['TRX chest press', 'Arnold press con kettlebell', 'Flexiones']),
            ('HOMBROS·TRÍCEPS', ['Extensión de tríceps con kettlebell', 'Elevaciones de hombro adelante', 'Banda triceps pushdown']),
            ('CORE·PIERNAS', ['Abs complex', 'Dead bug', 'Standing weighted hip abduction']),
            ('HIIT', ['Estocada con salto', 'Saltos laterales', 'Sentadilla con salto']),
        ]
    },
    {
        'name': 'Casa Press C — Nat',
        'circuits': [
            ('PIERNAS·GLÚTEOS', ['Plié squat', 'Loop band abducción parada', 'Sentadilla sumo']),
            ('PECHO·HOMBROS', ['Banda press de pecho', 'Vuelos laterales en equilibrio con kettlebell', 'Floor press']),
            ('HOMBROS·TRÍCEPS', ['Fondos en banco', 'Press militar', 'Extensión de tríceps con banda']),
            ('CORE·PIERNAS', ['Plancha con elevación alternada', 'Hollow body', 'Puente de glúteos una pierna']),
            ('HIIT', ['Mountain climbers', 'Burpees', 'Jumping jacks']),
        ]
    },
    {
        'name': 'Casa Press D — Nat',
        'circuits': [
            ('PIERNAS·GLÚTEOS', ['Sentadilla búlgara', 'Abducción con pausa', 'Zancadas con kettlebell']),
            ('PECHO·HOMBROS', ['Flexiones inclinadas', 'Press militar', 'TRX archer press']),
            ('HOMBROS·TRÍCEPS', ['Extensión de tríceps con kettlebell', 'Elevaciones de hombro adelante', 'Banda triceps pushdown']),
            ('CORE·PIERNAS', ['Plancha estrella con peso', 'Bird-dog', 'Loop band abducción parada']),
            ('HIIT', ['Sentadilla con salto', 'Saltos laterales', 'Estocada con salto']),
        ]
    },
    {
        'name': 'Casa Press E — Nat',
        'circuits': [
            ('PIERNAS·GLÚTEOS', ['Hip thrust a una pierna', 'Sumo squat to calf raise', 'Split squat pulses']),
            ('PECHO·HOMBROS', ['Flexiones diamante', 'Arnold press con kettlebell', 'TRX chest press']),
            ('HOMBROS·TRÍCEPS', ['Fondos en banco', 'Vuelos laterales', 'Extensión de tríceps con banda']),
            ('CORE·PIERNAS', ['Ab wheel', 'Plancha con chaleco', 'Extensión de cadera en 4 puntos']),
            ('HIIT', ['Jumping jacks', 'Mountain climbers', 'Burpees']),
        ]
    },
    {
        'name': 'Casa Press F — Nat',
        'circuits': [
            ('PIERNAS·GLÚTEOS', ['Elevated side reaches', 'Plié dips', 'Déficit reverse lunge']),
            ('PECHO·HOMBROS', ['Banda press de pecho', 'Vuelos laterales', 'Flexiones']),
            ('HOMBROS·TRÍCEPS', ['Extensión de tríceps con kettlebell', 'Elevaciones de hombro adelante', 'Banda triceps pushdown']),
            ('CORE·PIERNAS', ['Abs complex', 'Dead bug', 'Banda lateral walk']),
            ('HIIT', ['Estocada con salto', 'Sentadilla con salto', 'Saltos laterales']),
        ]
    },
]

NAT_PULL = [
    {
        'name': 'Casa Pull A — Nat',
        'circuits': [
            ('PIERNAS·GLÚTEOS', ['Plié dips', 'Elevated side reaches', 'Sentadilla búlgara con TRX']),
            ('PIERNAS·CORE', ['Peso muerto a una pierna', 'Empuje de cadera', 'Plancha con chaleco']),
            ('ESPALDA·BÍCEPS', ['TRX row', 'Bíceps curl', 'TRX face pull']),
            ('CORE·TRÍCEPS', ['Plancha estrella con peso', 'Banda pull-apart', 'Tríceps alto en TRX']),
            ('HIIT', ['Tuck jumps', 'Bear crawl', 'Saltos laterales']),
        ]
    },
    {
        'name': 'Casa Pull B — Nat',
        'circuits': [
            ('PIERNAS·GLÚTEOS', ['Sumo squat to RDL', 'Calf raises', 'Sentadilla búlgara']),
            ('PIERNAS·CORE', ['Puente de glúteos una pierna', 'Sentadilla goblet', 'Abs complex']),
            ('ESPALDA·BÍCEPS', ['Remo', 'Bíceps curl martillo', 'TRX face pull']),
            ('CORE·TRÍCEPS', ['Pallof press', 'Fondos en banco', 'Remo alto en TRX']),
            ('HIIT', ['Sentadilla con salto', 'Mountain climbers', 'Jumping jacks']),
        ]
    },
    {
        'name': 'Casa Pull C — Nat',
        'circuits': [
            ('PIERNAS·GLÚTEOS', ['Split squat pulses', 'Standing weighted hip abduction', 'Step up']),
            ('PIERNAS·CORE', ['Zancadas con chaleco de peso', 'Elevated side leg lifts', 'Dead bug']),
            ('ESPALDA·BÍCEPS', ['Dominadas australianas', 'Bíceps curl con banda', 'TRX power pull']),
            ('CORE·TRÍCEPS', ['Ab wheel', 'Extensión de tríceps con banda', 'Bíceps curl en equilibrio con kettlebell']),
            ('HIIT', ['Burpees', 'Estocada con salto', 'Bear crawl']),
        ]
    },
    {
        'name': 'Casa Pull D — Nat',
        'circuits': [
            ('PIERNAS·GLÚTEOS', ['Sumo squat to RDL', 'Abducción con pausa', 'Hip thrust a una pierna']),
            ('PIERNAS·CORE', ['Empuje de cadera en piso con peso', 'Sentadilla sumo', 'Plancha con elevación alternada']),
            ('ESPALDA·BÍCEPS', ['Remo alto en TRX', 'Bíceps curl en TRX', 'TRX face pull']),
            ('CORE·TRÍCEPS', ['Hollow body', 'Fondos en banco con chaleco', 'Remo alto en TRX']),
            ('HIIT', ['Tuck jumps', 'Saltos laterales', 'Mountain climbers']),
        ]
    },
    {
        'name': 'Casa Pull E — Nat',
        'circuits': [
            ('PIERNAS·GLÚTEOS', ['Step up', 'Elevated side reaches', 'Loop band abducción parada']),
            ('PIERNAS·CORE', ['Peso muerto a una pierna', 'Extensión de cadera en 4 puntos', 'Plancha con chaleco']),
            ('ESPALDA·BÍCEPS', ['TRX row', 'Bíceps curl concentrado', 'Dominadas australianas']),
            ('CORE·TRÍCEPS', ['Plancha toque de hombro', 'Tríceps alto en TRX', 'Banda pull-apart']),
            ('HIIT', ['Burpees', 'Sentadilla con salto', 'Jumping jacks']),
        ]
    },
    {
        'name': 'Casa Pull F — Nat',
        'circuits': [
            ('PIERNAS·GLÚTEOS', ['Plié squat', 'Calf raises', 'Déficit reverse lunge']),
            ('PIERNAS·CORE', ['Puente de glúteos una pierna', 'Sentadilla sumo', 'Bird-dog']),
            ('ESPALDA·BÍCEPS', ['TRX power pull', 'Bíceps curl martillo', 'Remo']),
            ('CORE·TRÍCEPS', ['Copenhague', 'Extensión de tríceps con kettlebell', 'Bíceps curl con banda']),
            ('HIIT', ['Estocada con salto', 'Bear crawl', 'Saltos laterales']),
        ]
    },
]

# ── Similar names (same exercise, different name) ─────────────────────────────
SIMILAR_PAIRS = [
    ('Empuje de cadera', 'Empuje de cadera en piso con peso', 'Mismo movimiento; "en piso con peso" es la versión con kettlebell'),
    ('Puente de glúteos una pierna', 'Hip thrust a una pierna', 'Mismo patrón unilateral, piso vs banco'),
    ('Sentadilla sumo', 'Plié squat', 'Mismo movimiento con diferente nombre'),
    ('Sentadilla sumo', 'Sumo squat to calf raise', 'Sentadilla sumo + elevación de talón'),
    ('Sentadilla sumo', 'Sumo squat to RDL', 'Sentadilla sumo + peso muerto'),
    ('Zancadas con chaleco de peso', 'Zancadas con kettlebell', 'Mismo movimiento, diferente carga'),
    ('Fondos en banco', 'Fondos en banco con chaleco', 'Misma base + chaleco de peso'),
    ('Remo', 'Remo alto en TRX', 'Ambos son remo pero ángulo distinto'),
    ('TRX row', 'Remo alto en TRX', 'Mismo equipo, variación de agarre'),
    ('Bíceps curl', 'Bíceps curl martillo', 'Mismo movimiento, diferente agarre'),
    ('Bíceps curl en TRX', 'Bíceps curl en equilibrio con kettlebell', 'Ambos curl con componente de estabilidad'),
    ('Flexiones con chaleco', 'Flexiones inclinadas', 'Misma base, diferente intensificación'),
    ('Sentadilla búlgara', 'Sentadilla búlgara con TRX', 'Misma base, TRX para asistencia'),
    ('Split squat pulses', 'Sentadilla búlgara', 'Patrón similar de zancada estática'),
]

def count_exercises(routines_list):
    """Count exercise frequency across all routines in a list."""
    counter = Counter()
    for r in routines_list:
        for _, exercises in r['circuits']:
            for ex in exercises:
                counter[ex] += 1
    return counter

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
    })
    shading_elm.append(shading)

def add_routine_table(doc, routines, title, freq_counter, high_freq_threshold=3):
    """Add a comparison table for a set of routines."""
    doc.add_heading(title, level=1)

    # Frequency legend
    p = doc.add_paragraph()
    run = p.add_run('Leyenda: ')
    run.bold = True
    run = p.add_run('■ Naranja ')
    run.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
    p.add_run(f'= ejercicio repetido {high_freq_threshold}+ veces en el grupo | ')
    run = p.add_run('■ Azul ')
    run.font.color.rgb = RGBColor(0x34, 0x98, 0xDB)
    p.add_run('= ejercicio repetido 2 veces')

    # Get circuit group names from first routine
    circuit_groups = [cg for cg, _ in routines[0]['circuits']]
    num_routines = len(routines)

    # Create table: header row + 1 row per circuit group
    table = doc.add_table(rows=1 + len(circuit_groups), cols=1 + num_routines)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    hdr.cells[0].text = 'Circuito'
    hdr.cells[0].paragraphs[0].runs[0].bold = True
    for i, r in enumerate(routines):
        letter = r['name'].split(' ')[2]  # A, B, C...
        hdr.cells[i + 1].text = letter
        hdr.cells[i + 1].paragraphs[0].runs[0].bold = True
        hdr.cells[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for row_idx, group in enumerate(circuit_groups):
        row = table.rows[row_idx + 1]
        row.cells[0].text = group
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(8)

        for col_idx, routine in enumerate(routines):
            circuit_data = routine['circuits'][row_idx]
            exercises = circuit_data[1]
            cell = row.cells[col_idx + 1]
            cell.text = ''
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)

            for ex_idx, ex in enumerate(exercises):
                freq = freq_counter[ex]
                run = para.add_run(ex)
                run.font.size = Pt(7.5)
                if freq >= high_freq_threshold:
                    run.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
                    run.bold = True
                elif freq >= 2:
                    run.font.color.rgb = RGBColor(0x34, 0x98, 0xDB)
                if ex_idx < len(exercises) - 1:
                    run = para.add_run('\n')
                    run.font.size = Pt(7.5)

    doc.add_paragraph()  # spacer

def add_frequency_table(doc, freq_counter, title, min_count=2):
    """Add a frequency summary table."""
    doc.add_heading(f'Frecuencia de ejercicios — {title}', level=2)
    items = [(ex, count) for ex, count in freq_counter.most_common() if count >= min_count]
    if not items:
        doc.add_paragraph('No hay ejercicios repetidos.')
        return

    table = doc.add_table(rows=1 + len(items), cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0]
    hdr.cells[0].text = 'Ejercicio'
    hdr.cells[1].text = 'Veces'
    hdr.cells[2].text = 'Frecuencia'
    for c in hdr.cells:
        c.paragraphs[0].runs[0].bold = True

    for i, (ex, count) in enumerate(items):
        row = table.rows[i + 1]
        row.cells[0].text = ex
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].text = str(count)
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        bar = '█' * count
        row.cells[2].text = bar
        if count >= 4:
            set_cell_shading(row.cells[0], 'FFF3E0')

    doc.add_paragraph()

def add_similar_table(doc):
    """Add table of similar-named exercises."""
    doc.add_heading('Ejercicios con nombres similares (posibles duplicados)', level=1)
    table = doc.add_table(rows=1 + len(SIMILAR_PAIRS), cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0]
    hdr.cells[0].text = 'Ejercicio 1'
    hdr.cells[1].text = 'Ejercicio 2'
    hdr.cells[2].text = 'Relación'
    for c in hdr.cells:
        c.paragraphs[0].runs[0].bold = True

    for i, (ex1, ex2, rel) in enumerate(SIMILAR_PAIRS):
        row = table.rows[i + 1]
        row.cells[0].text = ex1
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].text = ex2
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[2].text = rel
        row.cells[2].paragraphs[0].runs[0].font.size = Pt(8)
        set_cell_shading(row.cells[0], 'E8F5E9')
        set_cell_shading(row.cells[1], 'E8F5E9')

# ── Generate document ─────────────────────────────────────────────────────────

doc = Document()

# Set narrow margins
for section in doc.sections:
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.page_width = Cm(29.7)  # A4 landscape
    section.page_height = Cm(21.0)

doc.add_heading('Rutinas CASA — Tabla Comparativa', level=0)
doc.add_paragraph('Comparación de todas las rutinas de casa para Lean y Nat. Los ejercicios más repetidos están resaltados.')

# Lean Press
lean_press_freq = count_exercises(LEAN_PRESS)
add_routine_table(doc, LEAN_PRESS, 'LEAN — PRESS (Casa)', lean_press_freq)
add_frequency_table(doc, lean_press_freq, 'Lean Press')

# Lean Pull
lean_pull_freq = count_exercises(LEAN_PULL)
add_routine_table(doc, LEAN_PULL, 'LEAN — PULL (Casa)', lean_pull_freq)
add_frequency_table(doc, lean_pull_freq, 'Lean Pull')

# Nat Press
nat_press_freq = count_exercises(NAT_PRESS)
add_routine_table(doc, NAT_PRESS, 'NAT — PRESS (Casa)', nat_press_freq)
add_frequency_table(doc, nat_press_freq, 'Nat Press')

# Nat Pull
nat_pull_freq = count_exercises(NAT_PULL)
add_routine_table(doc, NAT_PULL, 'NAT — PULL (Casa)', nat_pull_freq)
add_frequency_table(doc, nat_pull_freq, 'Nat Pull')

# Similar exercises
add_similar_table(doc)

output_path = '/home/user/trainer2/Rutinas_CASA_Comparativa.docx'
doc.save(output_path)
print(f'✓ Documento generado: {output_path}')
